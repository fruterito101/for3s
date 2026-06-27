"""Procesa archivos adjuntos (imágenes, PDF, Word, Excel) para que For3s los
"lea" (2026-06-18, cierre del pendiente "Multimodal").

For3s era solo-texto: si le mandabas una foto o un PDF, lo ignoraba. Ahora
puede entender adjuntos. Dos familias según cómo los procesa Claude:

  · NATIVOS de Claude (visión / lectura de documentos) → se mandan tal cual,
    en base64, como bloques de contenido. Claude los "ve"/"lee" directamente:
      - imágenes  (jpg, png, gif, webp)  → bloque {"type":"image"}
      - PDF                               → bloque {"type":"document"}
  · NO nativos → extraemos el texto aquí (con una librería) y se lo pasamos a
    Claude como texto plano dentro de un bloque normal:
      - Word  (.docx) → python-docx
      - Excel (.xlsx) → openpyxl

El AUDIO queda fuera por ahora (decisión de diseño): Claude no transcribe audio;
se añadirá con Whisper en un segundo paso.

Esto NO es el MCP (que es solo GitHub). Devuelve BLOQUES de contenido listos
para inyectar en el payload de la API de Messages (ver llm.py.complete()).
"""

from __future__ import annotations

import base64
import io

# Límites de seguridad. La API de Claude acepta imágenes/PDF hasta cierto tamaño;
# Telegram de su lado limita los archivos. Cortamos antes para no reventar memoria
# ni mandar payloads que la API rechace.
MAX_BYTES = 20 * 1024 * 1024  # 20 MB por archivo (límite duro)
MAX_TEXTO_EXTRAIDO = 30_000  # chars de texto de Word/Excel que pasamos
MAX_FILAS_EXCEL = 500  # no volcar hojas gigantes enteras

# Umbral "MUY grande" para PDF/imagen que van a Claude como base64 (el dueño
# 2026-06-18): un archivo pesado infla el payload → Claude tarda y httpx puede
# dar ReadTimeout aun con timeout de 180s. Si supera esto, avisamos HONESTO en
# vez de intentarlo y morir. ~8 MB de binario ≈ payload ya muy grande.
MAX_BYTES_NATIVO = 8 * 1024 * 1024  # PDF/imagen que se manda en base64

# MIME types de imagen que la API de Claude acepta como bloque image.
_IMAGENES_OK = {
    "image/jpeg": "image/jpeg",
    "image/jpg": "image/jpeg",
    "image/png": "image/png",
    "image/gif": "image/gif",
    "image/webp": "image/webp",
}


class ArchivoNoSoportado(Exception):
    """El tipo de archivo no se puede procesar (con un mensaje para el usuario)."""


def _es_imagen(mime: str, nombre: str) -> str | None:
    """Devuelve el media_type normalizado si es imagen soportada, o None."""
    if mime in _IMAGENES_OK:
        return _IMAGENES_OK[mime]
    bajo = nombre.lower()
    for ext, mt in (
        (".jpg", "image/jpeg"),
        (".jpeg", "image/jpeg"),
        (".png", "image/png"),
        (".gif", "image/gif"),
        (".webp", "image/webp"),
    ):
        if bajo.endswith(ext):
            return mt
    return None


def _es_pdf(mime: str, nombre: str) -> bool:
    return mime == "application/pdf" or nombre.lower().endswith(".pdf")


def _es_word(mime: str, nombre: str) -> bool:
    return (
        mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or nombre.lower().endswith(".docx")
    )


def _es_excel(mime: str, nombre: str) -> bool:
    return mime in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ) or nombre.lower().endswith((".xlsx", ".xlsm"))


def _texto_de_word(datos: bytes) -> str:
    """Extrae el texto de un .docx (párrafos + tablas) con python-docx."""
    from docx import Document  # import perezoso: solo si llega un Word

    doc = Document(io.BytesIO(datos))
    partes: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            partes.append(p.text)
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if any(celdas):
                partes.append(" | ".join(celdas))
    texto = "\n".join(partes).strip()
    return texto or "(El documento Word no tiene texto extraíble.)"


def _texto_de_excel(datos: bytes) -> str:
    """Extrae las celdas de un .xlsx (todas las hojas) con openpyxl, como texto
    tabular. read_only + data_only (valores, no fórmulas) para ir ligero."""
    from openpyxl import load_workbook  # import perezoso

    wb = load_workbook(io.BytesIO(datos), read_only=True, data_only=True)
    partes: list[str] = []
    for hoja in wb.worksheets:
        partes.append(f"### Hoja: {hoja.title}")
        n = 0
        for fila in hoja.iter_rows(values_only=True):
            celdas = ["" if c is None else str(c) for c in fila]
            if any(c.strip() for c in celdas):
                partes.append(" | ".join(celdas))
                n += 1
            if n >= MAX_FILAS_EXCEL:
                partes.append(f"… (hoja truncada a {MAX_FILAS_EXCEL} filas)")
                break
    wb.close()
    texto = "\n".join(partes).strip()
    return texto or "(El Excel está vacío.)"


def _exigir_no_gigante(datos: bytes, tipo: str) -> None:
    """Para PDF/imagen (que van a Claude en base64): si superan MAX_BYTES_NATIVO,
    avisa HONESTO en vez de intentarlo y morir por timeout (2026-06-18).
    Un payload base64 enorme hace que Claude tarde > el timeout de httpx."""
    if len(datos) > MAX_BYTES_NATIVO:
        mb = len(datos) / 1024 / 1024
        lim = MAX_BYTES_NATIVO // 1024 // 1024
        raise ArchivoNoSoportado(
            f"Ese {tipo} pesa {mb:.0f} MB y es demasiado grande para procesarlo "
            f"completo (más de {lim} MB tarda tanto que se corta la conexión). "
            f"Si es un {tipo}, mándame solo las páginas/parte que te importan, o "
            f"un extracto — así sí puedo leerlo bien."
        )


def procesar_adjunto(datos: bytes, *, nombre: str = "", mime: str = "") -> list[dict]:
    """Convierte un archivo en BLOQUES de contenido para la API de Messages.

    Devuelve una lista de bloques (dicts) listos para meter en content[]:
      - imagen → [{"type":"image","source":{...base64...}}]
      - PDF    → [{"type":"document","source":{...base64...}}]
      - Word/Excel → [{"type":"text","text": "<texto extraído>"}]

    Lanza ArchivoNoSoportado (con mensaje legible) si no se puede procesar.
    """
    if not datos:
        raise ArchivoNoSoportado("El archivo llegó vacío.")
    if len(datos) > MAX_BYTES:
        mb = len(datos) / 1024 / 1024
        raise ArchivoNoSoportado(
            f"El archivo pesa {mb:.0f} MB; el máximo que puedo procesar son "
            f"{MAX_BYTES // 1024 // 1024} MB."
        )

    media = _es_imagen(mime, nombre)
    if media:
        _exigir_no_gigante(datos, "imagen")
        b64 = base64.standard_b64encode(datos).decode("ascii")
        return [
            {
                "type": "image",
                "source": {"type": "base64", "media_type": media, "data": b64},
            }
        ]

    if _es_pdf(mime, nombre):
        _exigir_no_gigante(datos, "PDF")
        b64 = base64.standard_b64encode(datos).decode("ascii")
        return [
            {
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": b64},
            }
        ]

    if _es_word(mime, nombre):
        texto = _texto_de_word(datos)[:MAX_TEXTO_EXTRAIDO]
        return [{"type": "text", "text": f"[Documento Word «{nombre}»]\n\n{texto}"}]

    if _es_excel(mime, nombre):
        texto = _texto_de_excel(datos)[:MAX_TEXTO_EXTRAIDO]
        return [{"type": "text", "text": f"[Hoja de cálculo «{nombre}»]\n\n{texto}"}]

    raise ArchivoNoSoportado(
        f"No sé leer ese tipo de archivo ({mime or nombre or 'desconocido'}). "
        "Puedo con imágenes, PDF, Word (.docx) y Excel (.xlsx). El audio llega pronto."
    )


def descripcion_corta(nombre: str, mime: str) -> str:
    """Etiqueta humana del tipo de archivo (para logs / mensajes)."""
    if _es_imagen(mime, nombre):
        return "imagen"
    if _es_pdf(mime, nombre):
        return "PDF"
    if _es_word(mime, nombre):
        return "documento Word"
    if _es_excel(mime, nombre):
        return "hoja de cálculo Excel"
    return "archivo"
